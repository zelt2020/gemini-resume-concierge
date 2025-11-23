from bs4 import BeautifulSoup
import requests
from docx import Document
from io import BytesIO

def scrape_job_url(url: str) -> str:
    headers = {"User-Agent": "Mozilla/5.0"}
    html = requests.get(url, headers=headers, timeout=15).text
    soup = BeautifulSoup(html, 'html.parser')
    return soup.get_text(separator="\n")[:15000]

def generate_ats_docx(resume: str, job: str) -> bytes:
    doc = Document()
    doc.add_paragraph("ATS-OPTIMIZED RESUME")
    doc.add_paragraph(f"Job keywords integrated: {job[:200]}")
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
